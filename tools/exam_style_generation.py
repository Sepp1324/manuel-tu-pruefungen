from __future__ import annotations

import hashlib
import re
from dataclasses import dataclass
from html import escape
from pathlib import Path

import pdfplumber
from docx import Document


@dataclass
class SourceDoc:
    path: Path
    chapter: int
    title: str
    pages: list[str]


ADMIN_RE = re.compile(
    r"(?i)\b("
    r"studierende|matrikelnummer|studienkennzahl|angaben zur pruefung|angaben zur prüfung|"
    r"beschriften sie|antwortzettel|pruefungsbogen|prüfungsbogen|sitzplatz|"
    r"smartwatch|smartphone|taschenrechner|tiss|tuwel|sommersemester|"
    r"bachelor studium|msc chemie|katharina\.ehrmann|@tuwien|tu wien|"
    r"vorlesung 164\.211|vorlesung 164\.221|t\.?\s*konegger|raquel\s+de\s+oro|"
    r"r\.\s*de\s+oro|raumnummer|allgemeine informationen|und analytik|"
    r"was sie ab heute verstehen werden|inhaltsverzeichnis|überblick über|ueberblick ueber|"
    r"administrative|lernunterlagen|ullmann|baerns|jess et al|164\.211|164\.221|"
    r"inhalt unterteilt|lernziel|lernziele|technologie zu beschreiben|"
    r"verfahrensprinzipien und prozessbedingungen|wichtige vorgaben|"
    r"pruefungsrelevant|prüfungsrelevant|vorausschau"
    r")\b"
)

SOURCE_NOISE_RE = re.compile(
    r"(?i)\b("
    r"doi:|https?://|www\.|copyright|literatur|sources?:|quellen?:|quelle:|"
    r"bildquellen?:|bildquelle|wikimedia commons|"
    r"systematic review|meta-regression|company overview|global headquarters|"
    r"regional headquarters|patent applications|microsoft copilot|"
    r"guardian|kurier\.at|derstandard|der standard|welt\.de|lv-evaluierung|tiss|"
    r"bertau|offermanns|salesch|zement-taschenbuch|industrielle anorganische chemie"
    r")\b"
)

LECTURE_ADMIN_RE = re.compile(
    r"(?i)\b("
    r"vor-?\s*(?:&amp;|&|und)\s*nachbereitung|vorlesungseinheiten|"
    r"positive absolvierung|mindestens 50\s*%|digitale geraete|digitale geräte|"
    r"abgabe am ende|zugewiesenen sitzplatz|anmeldeformalitaeten|anmeldeformalitäten"
    r")\b"
)

QUESTION_RE = re.compile(r"\?\s*(?:\[\d|$)")
FORMULA_RE = re.compile(r"\b[A-Z][a-z]?\d*(?:[A-Z][a-z]?\d*)+\b|->|→|Δ|∆")
OCR_NOISE_RE = re.compile(
    r"(?i)\(cid:\d+\)|[\uf0b7\uf0fc\ufffd�]|"
    r"\b(?:elleuqdli\.?\s*b|elleuq|elleu\.\s*q|nelleu\.\s*q|snommo\.\s*c|"
    r"aidemiki\.\s*w|kehtoto\.\s*f|ehcstue\.\s*d)\b"
)
TABLE_FRAGMENT_RE = re.compile(
    r"(?i)\b("
    r"investitionskosten|hauptvorteile|hauptnachteile|perspektive|infrastrukturbedarf|"
    r"flexibilitaet\s*/|flexibilität\s*/|produkte\s+standard|technologie\s+und\s+innovation|"
    r"produkt\s+festes\s+eisen|hochofeninfrast|flüssiges\s+roheisen\s+hochofen|"
    r"diaphragma-\s*amalgam-\s*membran|vorteile\s+anforderungen|"
    r"nachteile\s+geringe|soleaufbereitung|werkstoff\s+masse|"
    r"weltproduktion\s+verschiedener\s+werkstoffe|"
    r"hauptmerkmale\s+vorteile\s+nachteile|vorratsbunker|frischlanze|"
    r"einspulvorr|heizelektroden|verwendung\s+o\s+-?\s*einblasung|"
    r"ownership of|take-home messages|company|market|portfolio|headquarters|"
    r"manufacturing sites|innovation centers|patents|r&d|cagr|billion|metric tons"
    r")\b"
)
ENGLISH_PHRASE_RE = re.compile(
    r"(?i)\b("
    r"take-home messages|if we look towards|technology and innovation|ownership of|"
    r"leading innovation|mechanical recycling|chemical recycling|plastics recycling paths|"
    r"dissolving wood pulp|paper pulp|challenge:|circular economy|open hearth furnace|"
    r"basic oxygen furnace|blast furnace|fluidized beds|global organization|"
    r"high-quality|food packaging|medical grades|external use|internal confidential|"
    r"cellulose fibers|degree of polymerization|down the rabbit-hole|"
    r"polyolefins player|assets supporting|would mean|are hydrolyzed|becomes solubilized|"
    r"oxygen depolarized cathode|gas hourly space velocity|phase diagrams|engl\.|"
    r"filtration/degassing|xanthation|dissolution|homogenization|spinning|"
    r"sulfuric acid|aftertreatment|exhaust air recovery|spinbath recovery|"
    r"deauration|wet spinning"
    r")\b"
)
ENGLISH_WORDS = {
    "the", "and", "with", "for", "from", "this", "that", "which", "between",
    "process", "reaction", "temperature", "pressure", "growth", "market",
    "company", "global", "headquarters", "applications", "properties",
    "overview", "example", "formed", "regional", "manufacturing", "technology",
    "internal", "external", "confidential", "ownership", "innovation", "leading",
    "challenge", "mechanical", "chemical", "plastics", "paths", "messages",
    "towards", "industries", "dissolving", "pulp", "paper", "wood", "quality",
    "are", "would", "mean", "subsequent", "degree", "material", "transferring",
    "fibers", "suspended", "dissolved", "water", "dewatered", "drying", "machine",
    "then", "regenerated", "long", "player", "globally", "assets", "supporting",
    "reliable", "supply", "consistent", "includes", "largest", "rabbit", "hole",
    "already", "wiki", "becomes", "solubilized", "through", "hydrolyzed", "digester",
    "sulfonation", "after", "before", "during", "following", "transfer", "toward",
}
GERMAN_SIGNAL_WORDS = {
    "der", "die", "das", "und", "oder", "mit", "wird", "werden", "durch",
    "bei", "zur", "zum", "aus", "von", "fuer", "für", "als", "eine", "einer",
    "eines", "nicht", "nach", "vor", "rohstoff", "verfahren", "prozess",
    "reaktion", "herstellung", "eigenschaften", "anwendung", "beispiel",
}

PROCESS_TERMS = (
    "prozess", "verfahren", "synthese", "herstellung", "reaktion", "polymerisation",
    "cracking", "reforming", "vergasung", "roest", "röst", "elektrolyse", "hochofen",
    "claus", "solvay", "haber", "bosch", "bayer", "aufschluss", "recycling",
)
STRUCTURE_TERMS = (
    "struktur", "formel", "aufbau", "wiederholeinheit", "bindung", "monomer",
    "polymer", "molekular", "kristall", "netzwerk", "funktionelle gruppe",
)
COMPARE_TERMS = (
    "unterschied", "vs", "versus", "vergleich", "arten", "formen", "einteilung",
    "klassifizieren", "kategorie", "typen",
)
APPLICATION_TERMS = (
    "anwendung", "produkt", "eigenschaft", "vorteil", "nachteil", "kritik",
    "problem", "umwelt", "recycling", "rohstoff", "quelle", "vorkommen",
)
DOMAIN_TERMS = (
    "erdöl", "erdoel", "erdgas", "kohle", "claus", "synthesegas", "methanol",
    "olefin", "seife", "emulsion", "fett", "öl", "oel", "stärke", "staerke",
    "cellulose", "papier", "farbstoff", "polymer", "radikal", "stufenwachstum",
    "kunststoff", "mikroplastik", "stahl", "eisen", "hochofen", "oxid",
    "reduktion", "kupfer", "aluminium", "stickstoff", "ammoniak", "chlor",
    "natron", "soda", "schwefel", "gips", "kalk", "zement", "glas", "keramik",
    "polyamid", "polykondensat", "kevlar", "pigment",
)
TOPIC_STOPWORDS = {
    "verfahren", "prozess", "reaktion", "eigenschaften", "anwendungen",
    "probleme", "definition", "prinzip", "details", "vertiefung", "einheit",
    "grosschemie", "großchemie", "chemie", "technologie", "kurzpunkt",
}


def clean_text(text: str) -> str:
    text = text.replace("\u00ad", "")
    text = text.replace("\uf0b7", "•").replace("", "•").replace("▪", "•")
    text = text.replace("\uf0fc", "•").replace("", "•").replace("", "•")
    text = re.sub(r"\(cid:\d+\)", "", text)
    text = text.replace("→", "->")
    text = re.sub(r"([a-zäöüß])([A-ZÄÖÜ])", r"\1. \2", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"(?i)chemische technologie[n]? (?:organischer|anorganischer) stoffe", "", text)
    text = re.sub(r"(?i)chemische technologie (?:organischer|anorganischer) stoffe", "", text)
    text = re.sub(r"(?i)technische universitaet wien|technische universität wien|tu wien", "", text)
    text = re.sub(r"(?i)institut fuer chemische technologien|institut für chemische technologien", "", text)
    return text.strip()


def normalize_german_terms(text: str) -> str:
    replacements = {
        "Basic Oxygen Furnace": "Sauerstoffaufblasverfahren",
        "basic oxygen furnace": "Sauerstoffaufblasverfahren",
        "Open hearth furnace": "Siemens-Martin-Ofen",
        "open hearth furnace": "Siemens-Martin-Ofen",
        "Blast furnace": "Hochofen",
        "blast furnace": "Hochofen",
        "Fluidized Beds": "Wirbelschichtreaktoren",
        "fluidized beds": "Wirbelschichtreaktoren",
        "Depolymerization": "Depolymerisation",
        "depolymerization": "Depolymerisation",
        "Mechanical recycling": "mechanisches Recycling",
        "Chemical recycling": "chemisches Recycling",
        "Plastics Recycling": "Kunststoffrecycling",
        "block diagram": "Blockschema",
        "process flow diagram": "Verfahrensfliessschema",
        "piping and instrumentation diagram": "Rohrleitungs- und Instrumentenfliessschema",
        "Dual-Pressure": "Zweistufen-Druckfuehrung",
        "dual-pressure": "Zweistufen-Druckfuehrung",
        "INDUCTION FURNACE IF": "Induktionsofen",
        "Induction furnace": "Induktionsofen",
        "End-of-Life Recycling Rate": "Recyclingrate am Lebensende",
        "End-of-Life": "Lebensende",
        "Eo. L": "Lebensende",
        "EoL": "Lebensende",
        "Phase Diagrams": "Phasendiagramme",
        "Ellingham Diagrams": "Ellingham-Diagramme",
        "Ellingham Diagram": "Ellingham-Diagramm",
        "Freien Gibbs Gas Energie": "freie Gibbs-Energie",
        "Gibbs Gas Energie": "Gibbs-Energie",
        "Open hearth steel": "Siemens-Martin-Verfahren",
        "HOCHOFENS": "Hochofen",
        "Hochofens": "Hochofen",
        "Schwefels": "Schwefel",
        "Global Warming Potential": "Treibhauspotenzial",
        "High-quality": "hochwertig",
        "Na. Wa. Ros": "Nachwachsende Rohstoffe",
        "Direktreduction": "Direktreduktion",
    }
    for src, dst in replacements.items():
        text = text.replace(src, dst)
    return text


def extract_pdf_pages(path: Path) -> list[str]:
    pages: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = clean_text(page.extract_text(x_tolerance=1, y_tolerance=3) or "")
            if text:
                pages.append(text)
    return pages


def extract_docx_pages(path: Path) -> list[str]:
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = clean_text("\n".join(parts))
    chunks: list[str] = []
    buf: list[str] = []
    for line in text.splitlines():
        if re.match(r"^\s*(?:\d+\.\s+)?(?:Fossile|Nachwachsende|Makromolekulare|Polymer|Einheit)\b", line):
            if buf:
                chunks.append("\n".join(buf))
                buf = []
        buf.append(line)
    if buf:
        chunks.append("\n".join(buf))
    return [c for c in chunks if len(c) > 120]


def extract_document_pages(path: Path) -> list[str]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_pdf_pages(path)
    if suffix == ".docx":
        return extract_docx_pages(path)
    return []


def plain_line(line: str) -> str:
    line = normalize_german_terms(line)
    line = re.sub(r"\((?:engl\.|english)\s+[^)]*\)", "", line, flags=re.I)
    line = re.sub(r"\s+", " ", line).strip(" -;:>")
    line = re.sub(r"^\s*[•▪◦‣■●]+\s*", "", line)
    line = re.sub(r"(?i)\bT\.\s*Konegger\s*[–-].*$", "", line)
    line = re.sub(r"(?i)\bR\.\s*De\s+Oro\s+Calderon.*$", "", line)
    line = re.sub(r"(?i)\bRaquel\s+de\s+Oro\s+Calderon.*$", "", line)
    line = re.sub(r"(?i)\bChem\.\s*Techn\.\s*Anorg\.\s*Stoffe.*$", "", line)
    line = re.sub(r"(?i)\bChemische\s+Technologie\s+anorganische\s+Stoffe.*$", "", line)
    line = re.sub(r"(?i)\b(?:sources?|quellen?|quelle|bildquellen?|bildquelle):?.*$", "", line)
    line = re.sub(r"(?i)\b(?:wikimedia commons|quelle\s*\(unbekannt\)).*$", "", line)
    line = re.sub(r"(?i)\b(?:offermanns|salesch|ragaert|yin,|demets|akhras|fischer|winnacker|tegeder|haubner|hummel).*$", "", line)
    line = re.sub(r"(?i)\b[A-ZÄÖÜ][A-Za-zÄÖÜäöüß-]+(?:\s+und\s+[a-zäöüß.]+){0,2},\s*20\d{2},\s*p\d+\b.*$", "", line)
    line = re.sub(r"^\s*[a-z]\.\s+", "", line)
    line = re.sub(r"\s*\[\d+(?:[.,]\d+)?\]\s*$", "", line)
    line = re.sub(r"\s*\(\d+(?:[.,]\d+)?\)\s*$", "", line)
    return line.strip(" -;:>")


def english_heavy(line: str) -> bool:
    low = line.lower()
    if ENGLISH_PHRASE_RE.search(line):
        return True
    tokens = re.findall(r"\b[a-z]{3,}\b", low)
    english = sum(1 for token in tokens if token in ENGLISH_WORDS)
    german = sum(1 for token in tokens if token in GERMAN_SIGNAL_WORDS)
    has_umlaut = bool(re.search(r"[äöüßÄÖÜ]", line))
    if english >= 3:
        return True
    if english >= 2 and german == 0 and not has_umlaut:
        return True
    return False


def table_fragment(line: str) -> bool:
    if TABLE_FRAGMENT_RE.search(line):
        return True
    if line.count("/") >= 3 or line.count("|") > 1:
        return True
    chunks = re.split(r"\s{2,}|\s+-\s+", line)
    if len([c for c in chunks if len(c.split()) <= 4]) >= 4:
        return True
    return False


def noisy(line: str) -> bool:
    low = line.lower()
    if len(line) < 8:
        return True
    if OCR_NOISE_RE.search(line) or english_heavy(line) or table_fragment(line):
        return True
    if re.fullmatch(r"[\d\s.,:/()%-]+", line):
        return True
    if ADMIN_RE.search(line) or SOURCE_NOISE_RE.search(line) or LECTURE_ADMIN_RE.search(line):
        return True
    if any(x in low for x in (
        "guest lecture", "for external use", "generated by", "slide ", "recap",
        "terminology & concepts", "terminology and concepts", "market estimated",
        "usd ", "billion", "biggest players", "global organization",
        "what we do", "diversified portfolio", "borouge", "borstar", "sclairtech",
        "r&d employees", "patents", "high-voltage transmission", "medical grades",
        "flexible food packaging", "essential to everyday life",
        "population growth", "purchasing power", "textile growth", "metric tons", "cagr",
        "inhalt unterteilt", "technologie zu beschreiben", "verfahrensprinzipien",
        "wichtige vorgaben", "prüfungsrelevant", "pruefungsrelevant", "vorausschau",
    )):
        return True
    if line.count("|") > 1:
        return True
    if len(re.findall(r"\d", line)) > max(18, len(line) * 0.28):
        return True
    if len(re.findall(r"\b[A-Za-zÄÖÜäöüß]\b", line)) > 10:
        return True
    return False


def answer_fact_ok(line: str, topic: str = "") -> bool:
    if noisy(line) or QUESTION_RE.search(line):
        return False
    if len(line) > 230 or len(line.split()) < 4:
        return False
    low = f"{topic} {line}".lower()
    line_low = line.lower()
    domain_signal = any(t in low for t in DOMAIN_TERMS + PROCESS_TERMS + STRUCTURE_TERMS + APPLICATION_TERMS)
    german_signal = any(t in line_low for t in GERMAN_SIGNAL_WORDS) or bool(re.search(r"[äöüßÄÖÜ]", line))
    if not domain_signal and not FORMULA_RE.search(line) and not german_signal:
        return False
    if len(re.findall(r"[A-ZÄÖÜ][a-zäöüß]+", line)) > max(10, len(line.split()) * 0.55):
        return False
    return True


def meaningful_lines(text: str) -> list[str]:
    raw: list[str] = []
    for line in clean_text(text).splitlines():
        line = plain_line(line)
        if not line or noisy(line):
            continue
        raw.append(line)
    merged: list[str] = []
    for line in raw:
        if (
            merged
            and len(merged[-1]) < 120
            and len(line) < 120
            and not merged[-1].endswith((".", "?", "!", ":"))
            and not line.startswith(("•", "-", "a.", "b.", "c."))
        ):
            merged[-1] = f"{merged[-1]} {line}"
        else:
            merged.append(line)
    return [line[:260].strip() for line in merged if not noisy(line)]


def source_is_exam_example(source: SourceDoc) -> bool:
    local = " ".join(part.lower() for part in [source.path.name, *[p.name for p in source.path.parents[:3]]])
    return "pr_fungsvorbereitung" in local or "beispielpr" in local or re.search(r"\bpr_fung[_-]\d", local) is not None


def title_from_lines(lines: list[str], fallback: str) -> tuple[str, list[str]]:
    if not lines:
        return fallback, []
    first = lines[0]
    if QUESTION_RE.search(first) or len(first) > 120:
        return fallback, lines
    if len(lines) > 1 and len(first) < 35 and not first.endswith(":") and not lines[1].startswith("•"):
        joined = f"{first}: {lines[1]}"
        if len(joined) < 120 and not QUESTION_RE.search(lines[1]):
            return joined, lines[2:]
    return first.rstrip(":"), lines[1:]


def heading_like(line: str) -> bool:
    if noisy(line) or QUESTION_RE.search(line) or len(line) > 95:
        return False
    if re.match(r"^\d+[.)]\s", line) or FORMULA_RE.search(line):
        return False
    words = line.split()
    low = line.lower()
    if len(words) > 9:
        return False
    if line.endswith(":"):
        return True
    if any(term in low for term in DOMAIN_TERMS + PROCESS_TERMS + STRUCTURE_TERMS):
        return True
    return bool(re.match(r"^[A-ZÄÖÜ][A-Za-zÄÖÜäöüß/-]+(?:\s+[A-ZÄÖÜ][A-Za-zÄÖÜäöüß/-]+){0,4}$", line))


def section_blocks(lines: list[str], fallback: str) -> list[tuple[str, list[str]]]:
    heading_indexes = [idx for idx, line in enumerate(lines) if heading_like(line)]
    blocks: list[tuple[str, list[str]]] = []
    if not heading_indexes:
        title, rest = title_from_lines(lines, fallback)
        return [(title, rest)]
    for pos, idx in enumerate(heading_indexes):
        end = heading_indexes[pos + 1] if pos + 1 < len(heading_indexes) else len(lines)
        facts = lines[idx + 1:end]
        if len(facts) < 3:
            facts = lines[idx + 1:min(len(lines), idx + 8)]
        if len(facts) >= 2:
            blocks.append((lines[idx].rstrip(":"), facts))
    if not blocks:
        title, rest = title_from_lines(lines, fallback)
        blocks.append((title, rest))
    return blocks


def detail_topic_for(topic: str, facts: list[str], index: int, allow_anchor: bool = False) -> str:
    anchor = facts[0] if facts else ""
    anchor = re.sub(r"\([^)]{20,}\)", "", anchor)
    anchor = re.split(r"[.;]", anchor, 1)[0]
    anchor = re.sub(r"\s+", " ", anchor).strip(" -:;>")
    if len(anchor) > 55:
        anchor = " ".join(anchor.split()[:7])
    if allow_anchor and 12 <= len(anchor) <= 55 and not noisy(anchor):
        return f"{topic}: {anchor}"
    return f"{topic}: Vertiefung {index}"


def searchable(text: str) -> str:
    text = text.lower()
    text = text.replace("ä", "ae").replace("ö", "oe").replace("ü", "ue").replace("ß", "ss")
    return re.sub(r"[^a-z0-9]+", " ", text)


def topic_keywords(topic: str) -> list[str]:
    out: list[str] = []
    for word in re.findall(r"[A-Za-zÄÖÜäöüß0-9/-]{4,}", topic):
        key = searchable(word).strip()
        if len(key) < 5 or key in TOPIC_STOPWORDS:
            continue
        if re.fullmatch(r"\d+", key):
            continue
        out.append(key)
    return out[:4]


def topic_supported(topic: str, facts: list[str]) -> bool:
    keywords = topic_keywords(topic)
    if not keywords:
        return True
    fact_text = searchable(" ".join(facts))
    return any(keyword in fact_text for keyword in keywords)


def inferred_topic_from_facts(facts: list[str], fallback: str) -> str:
    fact_text = " ".join(facts[:6])
    process_names = re.findall(
        r"\b([A-ZÄÖÜ][A-Za-zÄÖÜäöüß]+(?:-[A-ZÄÖÜ]?[A-Za-zÄÖÜäöüß]+)*(?:-Prozess|-Verfahren|verfahren|prozess))\b",
        fact_text,
    )
    if process_names:
        return process_names[0][:90].strip(" :-")
    matches = re.findall(
        r"\b[A-ZÄÖÜ][A-Za-zÄÖÜäöüß0-9/-]{4,45}(?:\s+[A-ZÄÖÜ][A-Za-zÄÖÜäöüß0-9/-]{3,45}){0,2}\b",
        fact_text,
    )
    useful = [
        match.strip(" :-")
        for match in matches
        if not noisy(match) and topic_keywords(match)
    ]
    if useful:
        useful.sort(
            key=lambda x: (
                any(term in searchable(x) for term in DOMAIN_TERMS + PROCESS_TERMS + STRUCTURE_TERMS),
                len(topic_keywords(x)),
                len(x),
            ),
            reverse=True,
        )
        return useful[0][:90].strip(" :-")
    return fallback


def normalize_topic(title: str, facts: list[str], fallback: str) -> str:
    topic = re.sub(r"^\d+\.\s*", "", title)
    topic = topic.split("•", 1)[0]
    topic = re.sub(r"(?i)^(fossile rohstoffe|nachwachsende rohstoffe|polymerchemie|makromolekulare chemie|grosschemie|großchemie)\s*:?\s*", "", topic)
    topic = re.sub(r"\s+", " ", topic)
    topic = topic.strip(" :-")
    topic = normalize_german_terms(topic)
    if (
        len(topic) < 8
        or topic.lower() in {"einleitung", "grundlagen", "tools", "population"}
        or topic == fallback
        or re.fullmatch(r"(?:VO\d+|Einheit \d+).*", topic)
    ):
        text = " ".join(facts[:4])
        matches = re.findall(r"\b[A-ZÄÖÜ][A-Za-zÄÖÜäöüß0-9/-]{4,45}\b", text)
        scored = sorted(matches, key=lambda x: (any(t in x.lower() for t in DOMAIN_TERMS), len(x)), reverse=True)
        topic = scored[0] if scored else fallback
    fact_text = " ".join(facts[:5])
    process_names = re.findall(
        r"\b([A-ZÄÖÜ][A-Za-zÄÖÜäöüß]+(?:-[A-ZÄÖÜ]?[A-Za-zÄÖÜäöüß]+)*(?:-Prozess|-Verfahren|verfahren|prozess))\b",
        fact_text,
    )
    if process_names and (len(topic.split()) <= 2 or not any(term in topic.lower() for term in PROCESS_TERMS)):
        topic = process_names[0]
    if not topic_supported(topic, facts):
        topic = inferred_topic_from_facts(facts, fallback)
    if topic.endswith("-basierte") and re.search(fr"\b{re.escape(topic)}\s+Verfahren\b", " ".join(facts)):
        topic = f"{topic} Verfahren"
    return topic[:90].strip(" :-")


def classify_card(topic: str, facts: list[str]) -> str:
    text = " ".join([topic, *facts]).lower()
    if any(t in text for t in PROCESS_TERMS):
        return "exam_process"
    if any(t in text for t in COMPARE_TERMS):
        return "exam_compare"
    if any(t in text for t in STRUCTURE_TERMS):
        return "exam_structure"
    if any(t in text for t in APPLICATION_TERMS):
        return "exam_application"
    return "exam_concept"


def question_for(topic: str, kind: str, facts: list[str]) -> str:
    text = " ".join(facts).lower()
    needs_formula = FORMULA_RE.search(" ".join(facts)) or any(x in text for x in ("struktur", "formel", "gleichung"))
    if kind == "exam_process":
        aspects = "Ausgangsstoffe, Prozessfuehrung, wichtige Bedingungen, Produkte und Zweck"
        if "temperatur" in text or "druck" in text:
            aspects = "Ausgangsstoffe, Druck/Temperatur, Reaktionsschritte, Produkte und Zweck"
        if needs_formula:
            aspects += " sowie relevante Reaktions- oder Strukturformeln"
        return f"Erlaeutern Sie {topic}. Gehen Sie auf {aspects} ein."
    if kind == "exam_compare":
        return f"Vergleichen Sie {topic}. Nennen Sie die wichtigen Arten/Merkmale, Unterschiede und Beispiele."
    if kind == "exam_structure":
        extra = " und verwenden Sie Strukturformeln, falls sie pruefungsrelevant sind" if needs_formula else ""
        return f"Beschreiben Sie den chemischen Aufbau von {topic}. Erklaeren Sie Strukturmerkmale, Bindungen/Funktionsgruppen{extra}."
    if kind == "exam_application":
        return f"Nennen und erklaeren Sie die wichtigsten Eigenschaften, Anwendungen oder Probleme von {topic}."
    return f"Erklaeren Sie {topic} pruefungsnah. Gehen Sie auf Definition, zentrale Merkmale und mindestens ein Beispiel ein."


def fact_quality(line: str) -> int:
    low = line.lower()
    score = min(len(line), 120)
    if QUESTION_RE.search(line):
        score -= 35
    if any(t in low for t in DOMAIN_TERMS):
        score += 55
    if any(t in low for t in PROCESS_TERMS + STRUCTURE_TERMS + APPLICATION_TERMS):
        score += 30
    if FORMULA_RE.search(line):
        score += 25
    if low.startswith(("was ", "welche ", "wie ", "wodurch ", "warum ")):
        score -= 20
    return score


def answer_html(
    facts: list[str],
    source: SourceDoc,
    topic: str,
    min_items: int = 3,
    max_items: int = 7,
    heading: str = "Pruefungsantwort",
) -> str:
    clean_facts = [
        (idx, normalize_german_terms(fact))
        for idx, fact in enumerate(facts)
        if answer_fact_ok(fact, topic)
    ]
    if len(clean_facts) < min_items:
        return ""
    best = sorted(clean_facts, key=lambda item: fact_quality(item[1]), reverse=True)[:max_items]
    best.sort(key=lambda item: item[0])
    items = "".join(f"<li>{escape(fact)}</li>" for _, fact in best)
    return (
        f"<b>{escape(heading)} zu {escape(topic)}:</b>"
        f"<ul>{items}</ul>"
        "<b>Beim Antworten aktiv abdecken:</b> Definition/Prinzip, Prozess oder Struktur, "
        "wichtige Bedingungen, Produkte/Beispiele und typische Begruendung."
        f"<br><br><span class='source'>Quelle: {escape(source.title)}</span>"
    )


def card_weight(kind: str, facts: list[str]) -> int:
    weight = {
        "exam_process": 5,
        "exam_structure": 4,
        "exam_compare": 4,
        "exam_application": 3,
        "exam_concept": 2,
    }.get(kind, 1)
    if any(FORMULA_RE.search(f) for f in facts):
        weight += 1
    if len(facts) >= 4:
        weight += 1
    return weight


def build_candidates(source: SourceDoc, module: str, id_prefix: str, chapter_name: str) -> list[dict]:
    if source_is_exam_example(source):
        return []
    candidates: list[dict] = []
    for page_no, page in enumerate(source.pages, start=1):
        lines = meaningful_lines(page)
        if len(lines) < 3:
            continue
        for section_no, (title, rest) in enumerate(section_blocks(lines, chapter_name), start=1):
            facts = [line for line in rest if not noisy(line)]
            facts = [line for line in facts if len(line.split()) >= 3 or FORMULA_RE.search(line)]
            answer_facts = [line for line in facts if not QUESTION_RE.search(line)]
            if len(answer_facts) < 2:
                continue
            topic = normalize_topic(title, facts, chapter_name)
            if noisy(topic) or any(x in topic.lower() for x in (
                "market", "recap", "terminology", "essential materials", "population",
                "überblick", "ueberblick", "administrative", "vorlesung 164",
                "wichtige vorgaben", "pruefungsrelevant", "prüfungsrelevant",
                "kaum im einsatz", "hauptmerkmale", "vorteile nachteile",
            )):
                continue
            if re.match(r"(?i)^(in|bei|mit|durch|aus)\s+\w+", topic) and topic.endswith("."):
                continue
            kind = classify_card(topic, facts)
            question = question_for(topic, kind, facts)
            answer = answer_html(facts, source, topic)
            if not answer:
                continue
            key = re.sub(r"\W+", "", f"{module}:{source.title}:{page_no}:{section_no}:{question}:{answer}".lower())
            digest = hashlib.sha1(key.encode("utf-8")).hexdigest()[:12]
            candidates.append({
                "id": f"{id_prefix}:{digest}",
                "module": module,
                "deck": "anki",
                "kap": source.chapter,
                "sub": f"VO{source.chapter}" if source.chapter < 11 else "PV",
                "subname": chapter_name,
                "source": source.title,
                "kind": kind,
                "q": escape(question),
                "a": answer,
                "weight": card_weight(kind, facts),
            })
            for detail_no, start in enumerate(range(4, min(len(facts), 18), 4), start=1):
                detail_facts = facts[start:start + 7]
                if len(detail_facts) < 3:
                    continue
                detail_topic = detail_topic_for(topic, detail_facts, detail_no)
                detail_question = question_for(detail_topic, "exam_application", detail_facts)
                detail_answer = answer_html(detail_facts, source, detail_topic)
                if not detail_answer:
                    continue
                detail_key = re.sub(
                    r"\W+",
                    "",
                    f"{module}:{source.title}:{page_no}:{section_no}:detail:{detail_no}:{detail_question}:{detail_answer}".lower(),
                )
                detail_digest = hashlib.sha1(detail_key.encode("utf-8")).hexdigest()[:12]
                candidates.append({
                    "id": f"{id_prefix}:{detail_digest}",
                    "module": module,
                    "deck": "anki",
                    "kap": source.chapter,
                    "sub": f"VO{source.chapter}" if source.chapter < 11 else "PV",
                    "subname": chapter_name,
                    "source": source.title,
                    "kind": "exam_application",
                    "q": escape(detail_question),
                    "a": detail_answer,
                    "weight": max(1, card_weight("exam_application", detail_facts) - 1),
                })
            for quick_no, start in enumerate(range(0, min(len(facts) - 1, 24), 3), start=1):
                quick_facts = facts[start:start + 3]
                if len(quick_facts) < 2:
                    continue
                quick_topic = detail_topic_for(topic, quick_facts, quick_no)
                quick_question = (
                    f"Erklaeren Sie kurz {quick_topic}. "
                    "Nennen Sie zwei bis drei pruefungsrelevante Punkte."
                )
                quick_answer = answer_html(
                    quick_facts,
                    source,
                    quick_topic,
                    min_items=2,
                    max_items=3,
                    heading="Kurzantwort",
                )
                if not quick_answer:
                    continue
                quick_key = re.sub(
                    r"\W+",
                    "",
                    f"{module}:{source.title}:{page_no}:{section_no}:quick:{quick_no}:{quick_question}:{quick_answer}".lower(),
                )
                quick_digest = hashlib.sha1(quick_key.encode("utf-8")).hexdigest()[:12]
                candidates.append({
                    "id": f"{id_prefix}:{quick_digest}",
                    "module": module,
                    "deck": "anki",
                    "kap": source.chapter,
                    "sub": f"VO{source.chapter}" if source.chapter < 11 else "PV",
                    "subname": chapter_name,
                    "source": source.title,
                    "kind": "exam_fact",
                    "q": escape(quick_question),
                    "a": quick_answer,
                    "weight": max(1, card_weight("exam_concept", quick_facts) - 2),
                })
    return candidates


def balanced_cards(candidates: list[dict], target: int) -> list[dict]:
    seen: set[str] = set()
    unique: list[dict] = []
    for card in sorted(candidates, key=lambda c: (-c["weight"], c["kap"], c["source"], c["q"])):
        fingerprint = re.sub(r"\W+", "", (card["q"] + card["a"]).lower())[:260]
        if fingerprint in seen:
            continue
        seen.add(fingerprint)
        unique.append(card)
    buckets: dict[int, list[dict]] = {}
    for card in unique:
        buckets.setdefault(card["kap"], []).append(card)
    out: list[dict] = []
    while len(out) < target and any(buckets.values()):
        for kap in sorted(buckets):
            if buckets[kap] and len(out) < target:
                out.append(buckets[kap].pop(0))
    for order, card in enumerate(out, start=1):
        card["order"] = order
    return out


def generate_exam_style_cards(
    sources: list[SourceDoc],
    chapters: dict[int, str],
    module: str,
    id_prefix: str,
    target: int,
) -> list[dict]:
    candidates: list[dict] = []
    for source in sources:
        candidates.extend(build_candidates(source, module, id_prefix, chapters[source.chapter]))
    return balanced_cards(candidates, target)
