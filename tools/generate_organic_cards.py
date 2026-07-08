from __future__ import annotations

import hashlib
import json
import re
from dataclasses import dataclass
from html import escape
from pathlib import Path

import pdfplumber
from docx import Document


ROOT = Path(__file__).resolve().parents[2]
SOURCE_DIR = ROOT / "sources" / "organic_chem"
OUT_PATH = ROOT / "organicsr" / "app" / "seed_data.json"
RAW_OUT = ROOT / "organicsr" / "app" / "source_text_index.json"

TARGET_CARDS = 850

CHAPTERS = {
    1: "VO1 Einfuehrung und fossile Rohstoffe",
    2: "VO2 Raffinerieprozesse",
    3: "VO3 Raffinerie-Gastvortrag",
    4: "VO4 Raffinerie, Midstream und Vokabeln",
    5: "VO5 Nachwachsende Rohstoffe, Fette und Oele",
    6: "VO6 Kohlenhydrate und Staerke",
    7: "VO7 Cellulose",
    8: "VO8 Papier, Farbstoffe und Cellulose-Vertiefung",
    9: "VO9 Polymerchemie und Polymerisation",
    10: "VO10 Kunststoffrecycling",
    11: "Pruefungsvorbereitung und Beispielaufgaben",
}


@dataclass
class SourceText:
    path: Path
    chapter: int
    title: str
    text: str


def clean_text(text: str) -> str:
    text = text.replace("\u00ad", "")
    text = text.replace("\uf0b7", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"(?i)chemische technologien organischer stoffe", "", text)
    text = re.sub(r"(?i)technische universitaet wien|tu wien", "", text)
    return text.strip()


def chapter_for(path: Path) -> int:
    name = path.name.lower()
    parent = " ".join(p.lower() for p in path.parts)
    m = re.search(r"\bvo\s*([0-9]{1,2})(?:\b|_)", name)
    if m:
        n = int(m.group(1))
        if n <= 4:
            return n
        return n
    if "pr_fung" in parent or "beispiel" in parent:
        return 11
    if "vokabelsammlung vo1" in name:
        return 1
    if "vokabelsammlung vo2" in name:
        return 2
    if "vokabelsammlung vo3" in name:
        return 3
    if "vo4 vokabel" in name:
        return 4
    if "vokabelsammlung vo5" in name:
        return 5
    if "vokabelsammlung vo6" in name:
        return 6
    if "vokabelsammlung vo 7" in name:
        return 7
    if "vokabelsammlung vo8" in name:
        return 8
    if "vokabelsammlung vo 9" in name:
        return 9
    return 11


def extract_pdf(path: Path) -> str:
    pages: list[str] = []
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            text = page.extract_text(x_tolerance=1, y_tolerance=3) or ""
            text = clean_text(text)
            if text:
                pages.append(text)
    return "\n\n".join(pages)


def extract_docx(path: Path) -> str:
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return clean_text("\n".join(parts))


def load_sources() -> list[SourceText]:
    out: list[SourceText] = []
    for path in sorted(SOURCE_DIR.rglob("*")):
        if path.suffix.lower() not in {".pdf", ".docx"}:
            continue
        try:
            text = extract_pdf(path) if path.suffix.lower() == ".pdf" else extract_docx(path)
        except Exception as exc:
            print(f"skip {path}: {exc}")
            continue
        if len(text) < 120:
            continue
        chapter = chapter_for(path)
        out.append(SourceText(path, chapter, path.stem, text))
    return out


def noisy(s: str) -> bool:
    if re.match(r"^\(?\d+[\).:\s]", s):
        return True
    if any(x in s.lower() for x in ["http", "www.", ".de)", "guardian", "wiwo", "mtpa", "trl "]):
        return True
    if "|" in s:
        return True
    if re.search(r"\[Seite|\bSeite\s+\d+\b", s, re.I):
        return True
    if len(re.findall(r"\d", s)) > max(10, len(s) * 0.18):
        return True
    if s.count("/") + s.count("|") + s.count("->") > 4:
        return True
    if len(re.findall(r"\b[a-zA-ZÄÖÜäöüß]\b", s)) > 8:
        return True
    return False


def sentences(text: str) -> list[str]:
    raw_lines = []
    for line in text.splitlines():
        line = re.sub(r"\s+", " ", line).strip(" -;:")
        if not line or re.fullmatch(r"\d{1,3}", line):
            continue
        if re.search(r"^(seite|slide|vorlesung|universit|institut)\b", line, re.I):
            continue
        raw_lines.append(line)
    chunks = []
    for line in raw_lines:
        chunks.extend(re.split(r"(?<=[.!?])\s+|(?:\s+•\s+)", line))
    cleaned: list[str] = []
    for s in chunks:
        s = re.sub(r"\s+", " ", s).strip(" -;:")
        if 55 <= len(s) <= 360 and len(s.split()) >= 7:
            if not re.search(r"[A-Za-zÄÖÜäöüß]", s):
                continue
            if noisy(s):
                continue
            cleaned.append(s)
    return cleaned


TERM_RE = re.compile(
    r"\b([A-ZÄÖÜ][A-Za-zÄÖÜäöüß0-9][A-Za-zÄÖÜäöüß0-9 /()\\-]{3,55})\b"
)
GENERIC_TERMS = {
    "anwendung", "alternativ", "location", "beispiel", "vorteile", "nachteile",
    "eigenschaften", "produkte", "prozess", "verfahren", "quelle",
}

SOURCE_NOISE_RE = re.compile(
    r"(?i)\b(systematic review|meta-regression|doi:|trademark|trademarks|"
    r"company overview|company information|internal: confidential|global headquarters|"
    r"regional headquarters|patent applications|workforce|microsoft copilot|"
    r"marktresearchfuture|guardian|kurier\.at|derstandard|welt\.de|"
    r"lv-evaluierung|tiss)\b"
)
COMPANY_RE = re.compile(
    r"\b[A-ZÄÖÜ][A-Za-zÄÖÜäöüß'.-]+\s+(AG|GmbH|Inc|Ltd|LLC|International)\b"
)
AUTHOR_RE = re.compile(
    r"\b[A-ZÄÖÜ][A-Za-zÄÖÜäöüß'.-]+\s+[A-Z]\.?\s+[A-ZÄÖÜ][A-Za-zÄÖÜäöüß'.-]+\b"
    r"|\b[A-ZÄÖÜ][A-Za-zÄÖÜäöüß'.-]+\s+[A-ZÄÖÜ][A-Za-zÄÖÜäöüß'.-]+,\s+"
    r"[A-ZÄÖÜ][A-Za-zÄÖÜäöüß'.-]+\s+[A-ZÄÖÜ][A-Za-zÄÖÜäöüß'.-]+\b"
)
LECTURE_ADMIN_RE = re.compile(
    r"(?i)\b(vor-?\s*(?:&amp;|&|und)\s*nachbereitung|vorlesungseinheiten|"
    r"prüfungsbogen|pruefungsbogen|prüfungsbeginn|pruefungsbeginn|"
    r"prüfungsantritt|pruefungsantritt|angabezettel|antwortzettel|"
    r"abgabe am ende|zugewiesenen sitzplatz|digitale geräte|digitale geraete|"
    r"sammelort|anmeldeformalitäten|anmeldeformalitaeten|"
    r"positive absolvierung|mindestens 50% der punkte|vorlesungsteile)\b"
)


def valid_term(term: str) -> bool:
    t = term.strip(" ,.;:-")
    low = t.lower()
    if low in GENERIC_TERMS:
        return False
    if t.count("(") != t.count(")"):
        return False
    if re.search(r"[A-Za-zÄÖÜäöüß]{14,}", t):
        return False
    if len(t.split()) > 6:
        return False
    if noisy(t):
        return False
    return True


def irrelevant_source_fact(text: str) -> bool:
    if SOURCE_NOISE_RE.search(text):
        return True
    if COMPANY_RE.search(text):
        return True
    if AUTHOR_RE.search(text) and re.search(r"(?i)\b(review|meta-regression|study|journal|doi)\b", text):
        return True
    if re.search(r"(?i)\b(location|headquarters|corporate hub|patent|source:|notes?:)\b", text):
        return True
    if LECTURE_ADMIN_RE.search(text):
        return True
    return False


def pick_term(sentence: str) -> str | None:
    candidates = []
    for m in TERM_RE.finditer(sentence):
        term = m.group(1).strip(" ,.;:-")
        if len(term) < 4 or not valid_term(term):
            continue
        if term.lower().startswith(("seite", "abbildung", "quelle", "beispiel")):
            continue
        score = len(term) + (18 if any(c.isupper() for c in term[1:]) else 0)
        if any(x in term.lower() for x in ["poly", "cellulose", "raffin", "crack", "destillation", "kohlen", "fett", "oel", "recycling"]):
            score += 22
        candidates.append((score, term))
    if not candidates:
        return None
    return max(candidates)[1]


def answer_html(sentence: str, source: SourceText) -> str:
    return (
        f"{escape(sentence)}"
        f"<br><br><span class='source'>Quelle: {escape(source.title)} ({CHAPTERS[source.chapter]})</span>"
    )


def make_cards(sources: list[SourceText]) -> list[dict]:
    cards: list[dict] = []
    seen: set[str] = set()

    def add(source: SourceText, question: str, answer: str, kind: str, weight: int = 1):
        key = re.sub(r"\W+", "", (question + answer).lower())
        digest = hashlib.sha1(key.encode("utf-8")).hexdigest()[:12]
        if digest in seen:
            return
        seen.add(digest)
        cards.append({
            "id": f"org:{digest}",
            "deck": "anki",
            "kap": source.chapter,
            "sub": f"VO{source.chapter}" if source.chapter < 11 else "PV",
            "subname": CHAPTERS[source.chapter],
            "source": source.title,
            "kind": kind,
            "q": question,
            "a": answer,
            "weight": weight,
        })

    for source in sources:
        for s in sentences(source.text):
            if irrelevant_source_fact(s):
                continue
            term = pick_term(s)
            if term:
                cloze = re.sub(re.escape(term), "_____", s, count=1)
                add(
                    source,
                    f"Ergaenze den zentralen Begriff: {escape(cloze)}",
                    answer_html(f"{term} - {s}", source),
                    "cloze",
                    2,
                )
            lead = term or "diesem Thema"
            if re.search(r"\b(ist|sind|wird|werden|entsteht|erfolgt|dient|besteht|enthaelt|kann|koennen|fuehrt|fuehren)\b", s, re.I):
                add(
                    source,
                    f"Was musst du zu {escape(lead)} wissen?",
                    answer_html(s, source),
                    "concept",
                    2,
                )
            if re.search(r"\b(Vorteil|Nachteil|Anwendung|Beispiel|Verfahren|Prozess|Reaktion|Produkt|Eigenschaft)\b", s, re.I):
                add(
                    source,
                    f"Nenne den relevanten Punkt zu {escape(lead)}.",
                    answer_html(s, source),
                    "exam_fact",
                    1,
                )

    if len(cards) < 620:
        for source in sources:
            clean_units = [s for s in sentences(source.text) if not noisy(s)]
            if not clean_units:
                continue
            for idx, unit in enumerate(clean_units[:4], start=1):
                if irrelevant_source_fact(unit):
                    continue
                add(
                    source,
                    f"Welche Kernaussage steht in {escape(source.title)} zu diesem Punkt?",
                    answer_html(unit, source),
                    "source_fact",
                    1,
                )
                if len(cards) >= 620:
                    break
            if len(cards) >= 620:
                break

    # Keep a balanced set across chapters and card types.
    cards.sort(key=lambda c: (-c["weight"], c["kap"], c["kind"], c["q"]))
    buckets: dict[int, list[dict]] = {}
    for card in cards:
        buckets.setdefault(card["kap"], []).append(card)
    balanced: list[dict] = []
    while len(balanced) < TARGET_CARDS and any(buckets.values()):
        for kap in sorted(buckets):
            if buckets[kap] and len(balanced) < TARGET_CARDS:
                balanced.append(buckets[kap].pop(0))
    for i, card in enumerate(balanced, start=1):
        card["order"] = i
    return balanced


def main() -> None:
    sources = load_sources()
    cards = make_cards(sources)
    coverage = []
    for kap, title in CHAPTERS.items():
        coverage.append({
            "kap": kap,
            "name": title,
            "anki": sum(1 for c in cards if c["kap"] == kap),
            "sources": sorted({s.title for s in sources if s.chapter == kap}),
        })
    payload = {
        "title": "Chemische Technologien Organischer Stoffe",
        "exam_date": "2026-09-21",
        "chapters": {str(k): v for k, v in CHAPTERS.items()},
        "cards": cards,
        "coverage": coverage,
        "source_count": len(sources),
    }
    OUT_PATH.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    RAW_OUT.write_text(json.dumps([
        {"path": str(s.path.relative_to(ROOT)), "chapter": s.chapter, "title": s.title, "chars": len(s.text), "text": s.text[:4000]}
        for s in sources
    ], ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"wrote {len(cards)} cards from {len(sources)} sources to {OUT_PATH}")
    for row in coverage:
        print(f"{row['kap']:>2}: {row['anki']:>3} {row['name']}")


if __name__ == "__main__":
    main()
